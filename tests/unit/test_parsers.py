"""Document parsing — generates real docx/xlsx files and parses them."""

import io

from etki.extraction.parsers import parse_document


def test_parse_plain_text_filters_short_lines():
    data = "rapora filtre eklensin\nSSO entegrasyonu yapılsın\nab\n".encode()
    _, items = parse_document("notlar.txt", data)
    assert "rapora filtre eklensin" in items
    assert "ab" not in items  # short/meaningless lines are filtered out


def test_parse_csv_rows():
    data = "id,talep\n1,rapora filtre eklensin\n2,kripto ödeme eklensin\n".encode()
    _, items = parse_document("crler.csv", data)
    assert any("rapora filtre" in it for it in items)
    assert any("kripto" in it for it in items)


def test_parse_docx_paragraphs():
    from docx import Document

    document = Document()
    document.add_paragraph("rapora filtre eklensin")
    document.add_paragraph("SSO entegrasyonu yapılsın")
    buffer = io.BytesIO()
    document.save(buffer)
    _, items = parse_document("talepler.docx", buffer.getvalue())
    assert len(items) == 2
    assert "SSO entegrasyonu yapılsın" in items


def test_parse_xlsx_rows():
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["Talep"])
    sheet.append(["rapora filtre eklensin"])
    sheet.append(["kripto para ile ödeme eklensin"])
    buffer = io.BytesIO()
    workbook.save(buffer)
    _, items = parse_document("crler.xlsx", buffer.getvalue())
    assert any("rapora filtre" in it for it in items)
    assert any("kripto" in it for it in items)


def test_cp1254_turkish_upload_keeps_exclusion_polarity():
    """W3: a Windows-1254 export used to decode with errors='replace' — 'hariçtir'
    became 'hari�tir', the keyword vanished and the clause polarity INVERTED."""
    import asyncio

    from etki.extraction.scope_extractor import HeuristicScopeExtractor

    contract = "## Madde 7.1 — Kapsam\nSSO entegrasyonu hariçtir.\n"
    text, _ = parse_document("sozlesme.txt", contract.encode("cp1254"))
    assert "hariçtir" in text  # decoded via the CP1254 fallback, not replaced
    items = asyncio.run(HeuristicScopeExtractor().extract("C", text))
    assert items[0].polarity.value == "EXCLUDED"


def test_undecodable_text_is_refused_loudly():
    import pytest
    from etki.extraction.parsers import DocumentUnreadable

    with pytest.raises(DocumentUnreadable):
        parse_document("garip.txt", b"\xff\xfe\x00\x01\x81\x91")


def test_docx_preserves_document_order_and_bullet_markers(tmp_path):
    """W3: tables used to be appended AFTER all paragraphs (attributed to the
    last clause) and Word list bullets lost their markers (never matched
    _BULLET). Both fixed by walking the body in document order."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("Madde 3 — Entegrasyon")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "ERP sistemi"
    table.rows[0].cells[1].text = "dahildir"
    doc.add_paragraph("Madde 7 — Hariç Tutulanlar")
    doc.add_paragraph("SSO entegrasyonu kapsam dışıdır.", style="List Bullet")
    path = tmp_path / "sozlesme.docx"
    doc.save(path)

    text, _ = parse_document("sozlesme.docx", path.read_bytes())
    lines = text.splitlines()
    assert lines.index("ERP sistemi dahildir") < lines.index("Madde 7 — Hariç Tutulanlar")
    assert "- SSO entegrasyonu kapsam dışıdır." in lines  # bullet marker restored
