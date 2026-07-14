"""Converts uploaded documents (txt/md/csv/docx/xlsx/pdf) to text + request lines.

`parse_document(filename, data) -> (full_text, items)`:
  - full_text: the full text, for scope extraction.
  - items: a list of meaningful lines/paragraphs/rows for triage (each a request candidate).
Heavy libraries (docx/openpyxl/pypdf) are lazily imported only for the relevant format.
"""

from __future__ import annotations

import csv
import io
from pathlib import Path


def _meaningful(line: str) -> bool:
    stripped = line.strip()
    return len(stripped) > 3 and any(c.isalpha() for c in stripped)


def _from_text(text: str) -> tuple[str, list[str]]:
    items = [ln.strip() for ln in text.splitlines() if _meaningful(ln)]
    return text, items


def _from_csv(data: bytes, delimiter: str) -> tuple[str, list[str]]:
    text = data.decode("utf-8", errors="replace")
    items: list[str] = []
    for row in csv.reader(io.StringIO(text), delimiter=delimiter):
        joined = " ".join(cell.strip() for cell in row if cell.strip())
        if _meaningful(joined):
            items.append(joined)
    return text, items


def _from_docx(data: bytes) -> tuple[str, list[str]]:
    from docx import Document

    document = Document(io.BytesIO(data))
    items = [p.text.strip() for p in document.paragraphs if _meaningful(p.text)]
    for table in document.tables:
        for row in table.rows:
            joined = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if _meaningful(joined):
                items.append(joined)
    return "\n".join(items), items


def _from_xlsx(data: bytes) -> tuple[str, list[str]]:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    items: list[str] = []
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            joined = " ".join(cells)
            if _meaningful(joined):
                items.append(joined)
    workbook.close()
    return "\n".join(items), items


def _from_pdf(data: bytes) -> tuple[str, list[str]]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    return _from_text(text)


def parse_document(filename: str, data: bytes) -> tuple[str, list[str]]:
    ext = Path(filename).suffix.lower()
    if ext == ".csv":
        return _from_csv(data, ",")
    if ext == ".tsv":
        return _from_csv(data, "\t")
    if ext == ".docx":
        return _from_docx(data)
    if ext == ".xlsx":
        return _from_xlsx(data)
    if ext == ".pdf":
        return _from_pdf(data)
    # .txt / .md / unknown → plain text
    return _from_text(data.decode("utf-8", errors="replace"))
