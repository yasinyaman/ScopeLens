"""Builds a client-/dispute-facing Word (.docx) report from a case file (python-docx).

For a non-technical user: decisions + reasoning + cited contract clauses + effort +
risk + audit trail in one document — sendable to the client / usable as a dispute file.
"""

from __future__ import annotations

import io

from etki.core.models import AuditEvent, CaseFile


def build_case_report(case: CaseFile, audit: list[AuditEvent]) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading(f"Karar Raporu — {case.request_id}", level=0)
    document.add_paragraph(
        f"Proje: {case.project_id or '—'}  ·  Durum: {case.status.value}"
    )
    document.add_paragraph(f"Talep: {case.raw_request}")
    document.add_paragraph(
        "Bu rapor bir karar destek önerisidir; nihai karar PMO'dadır (copilot, autopilot değil)."
    )

    for i, decision in enumerate(case.decisions):
        item = case.sub_requests[i].item if i < len(case.sub_requests) else decision.request_id
        document.add_heading(f"{i + 1}. {item}", level=1)
        document.add_paragraph(
            f"Karar: {decision.decision.value}  (güven %{decision.confidence * 100:.0f})"
        )
        est = decision.effort_estimate
        document.add_paragraph(f"Efor: {est.low}–{est.high} {est.unit}  ({est.basis})")
        risk_line = f"Risk: {decision.risk.level.value}"
        if decision.risk.escalation:
            risk_line += "  — 24 saatte eskalasyon"
        document.add_paragraph(risk_line)
        document.add_paragraph(f"Gerekçe: {decision.evidence.reasoning}")
        if decision.evidence.impacted_modules:
            document.add_paragraph(
                "Etkilenen modüller: " + ", ".join(decision.evidence.impacted_modules)
            )
        if decision.evidence.cited_clauses:
            document.add_paragraph("Atıf yapılan sözleşme maddeleri:")
            for clause in decision.evidence.cited_clauses:
                document.add_paragraph(
                    f"{clause.source_clause or clause.id} "
                    f"[{clause.polarity.value}]: {clause.description}",
                    style="List Bullet",
                )
        if decision.cr_draft is not None:
            document.add_paragraph(f"CR yayılım analizi: {decision.cr_draft.impact_analysis}")
            document.add_paragraph(f"Tahmini efor: {decision.cr_draft.cost}")

    if audit:
        document.add_heading("Denetim İzi (yeniden kurgulanabilir)", level=1)
        for event in audit:
            document.add_paragraph(
                f"#{event.seq} [{event.actor}] {event.action} — {event.detail}",
                style="List Number",
            )

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
